from docx import Document
import re

def find_par(document, text):
    num_par = 0
    for paragraph in document.paragraphs:
        if (text in paragraph.text.strip()):
            if ((len(paragraph.text.strip()) - len(text)) == 0):
                return num_par
        num_par = num_par + 1
    return -1

def find_biblio():
    document = Document("d:/Сибгути/1 курс/Экономика/Курсовая Дефицит бюджета (ИИ-051 Носков Кирилл).docx") #изменить путь в передаваемую переменную

    i = 0
    text = 'Список литературы'
    text2 = 'Список использованных источников'

    first = find_par(document, text)
    second = find_par(document, text2)
    maxi = max(first, second)
    if maxi == -1:
        return 'Отсутствуют источники или их заголовок задан неверно' #прописать сообщение об ошибке (отсутствие источников)
    else:
        k = ''
        flg = 0
        number = 1
        while i < len(document.paragraphs):
            if 'Приложение' in document.paragraphs[i].text:
                flg = 0
            if flg == 1:
                if re.sub(r"^\s+", "", document.paragraphs[i].text) != '':
                    k = k+str(number)+'. '+document.paragraphs[i].text+'\n'
                    number = number + 1
            if i == maxi:
                flg = 1
            i = i + 1
        return k 

####### строки ниже только для проверки!!!
l = find_biblio() 
print(l)   