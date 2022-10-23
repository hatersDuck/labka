# 5.1.2 На титульном листе приводят следующие сведения:
# а) наименование министерства (ведомства) или другого структурного образования, в систему которого входит
# организация-исполнитель;
# б) наименование (полное и сокращенное) организации - исполнителя НИР;
# в) индекс Универсальной десятичной классификации (УДК) по ГОСТ 7.90;
# г) номера, идентифицирующие отчет:

# Цвет шрифта должен быть черным, размер шрифта - не менее 12 пт. Рекомендуемый тип шрифта
# для основного текста отчета - Times New Roman. Полужирный шрифт применяют только для заголовков разделов и подразделов,
# заголовков структурных элементов. Использование курсива допускается для обозначения объектов (биология, геология,
# медицина, нанотехнологии, генная инженерия и др.) и написания терминов (например, in vivo, in vitro) и иных объектов и
# терминов на латыни.

# Текст отчета следует печатать, соблюдая следующие размеры полей: левое - 30 мм, правое - 15 мм, верхнее и нижнее - 20
# мм. Абзацный отступ должен быть одинаковым по всему тексту отчета и равен 1,25 см.

import string
from docx import Document
import docx2txt
import re
from biblio import checkBiblio

class checkDocument():
    def __Init__(self, way) -> None:

        """

        errosTitle - Ошибки в оформлении работы
        studentInfo - Информация о студенте
        document - Работа студента
        biblio - Список литературы

        """
        self.way = way

        self.document = Document(self.way)
        self.strf = docx2txt.process(self.way).splitlines()

        self.__titleInfo__()
        self.__studentInfo__()
        self.__biblioInfo__()



    def __str__(self):
        return "\n".join(self.errorsTitle)

    def __titleInfo__(self):
        base_errors = [0,0,0,0,0] #содержание, введение, заключение, список литературы
        text_errors = ['содержание', 'введение', 'заключение', 'список литературы', 'cписок использованных источников']
        zagl = []
        self.errorsTitle = []
        i = 0
        flg = 0
        k = 0
        for i in range(len(self.strf)):
            txt = re.sub("[0-9]+|[\.]", "", self.strf[i].lower()).strip()
            j = 0
            for j in range(len(text_errors)):
                if (txt == text_errors[j]):
                    base_errors[j]+=1
                    flg = 1
            
            if (txt in zagl):
                k += 1
            if (base_errors[1] == 2):
                flg = 2
            if (flg == 1):
                zagl.append(txt) if txt else zagl
        flg = 0
        if (k+1 != len(zagl)):
            flg = 1   

        i = 0
        for i in range(len(base_errors)-1): #### ошибки ключевых заголовков
            if (i == 0):
                if (base_errors[i] == 0):
                    flg = 1
            elif (i < 3):
                if (base_errors[i] < 2):
                    flg = 1
            else:
                if (base_errors[i] != 2) and (base_errors[i+1] != 2):
                    flg = 1
        if (flg == 1):
            self.errorsTitle.append('Ошибка в заголовках или содержании')

        sections = self.document.sections
        #len(sections)
        flg = 1
        s = ['','']
        for section in sections:
            if ((round(section.top_margin.cm, 2) != 2) or (round(section.bottom_margin.cm, 2) != 2) 
            or (round(section.left_margin.cm, 2) != 3) or (round(section.right_margin.cm, 2) != 1)):
                flg = 0
            
            if (section.footer.is_linked_to_previous == True):
                s[0] = 1
            if (section.footer.is_linked_to_previous != True):
                s[1] = 1
            
        if flg == 0:
            self.errorsTitle.append('Неправильные поля')
        if (s[0] != 1 or s[1] != 1):
            self.errorsTitle.append('Ошибка в нумерации')
        
        flg = 0
        for paragraph in self.document.paragraphs:
            a = 0
            b = 0
            if flg == 1:
                for r in paragraph.runs:
                    if (r.font.size.pt != 14):
                        a = 1
                    if (r.font.name != None) and (r.font.name != 'Times New Roman'):
                        b = 1

                if (paragraph.paragraph_format.line_spacing != 1.5):
                    self.errorsTitle.append('Неверный межстрочный интервал')
                if (a == 1):
                    self.errorsTitle.append('Неверный размер шрифта')
                if (b == 1):
                    self.errorsTitle.append('Неверный шрифт') 
                break
            if (paragraph.text.strip().lower() == 'введение'):
                flg = 1

    def __studentInfo__(self):
        text = ['Выполнил:', 'Группа:', 'гр.', 'введение']
        i = 0
        self.studentInfo = []
        for i in range(len(self.strf)):
            txt = self.strf[i]
            if (text[0] in txt):
                v = (re.sub(f'{text[0]}', "", self.strf[i]).strip()).split(' ')
                v = v[0]+' '+v[1][0]+'. '+v[2][0]+'.'
                self.studentInfo.append(v)
            elif (text[1] in txt):
                v = re.sub(f'{text[1]}+|{text[2]}', "", self.strf[i]).strip()
                self.studentInfo.append(v)
            elif (text[2] in txt):
                v = txt.partition(f'{text[2]}')[2].strip()
                self.studentInfo.append(v)
            if (txt.lower().split == text[3]):
                break

    def __biblioInfo__(self):
        i = 0

        maxi = max(
            findParagraph('Список литературы', self.document), 
            findParagraph('Список использованных источников', self.document),
            findParagraph('Библиографический список', self.document)
        )
        if maxi == -1:
            self.biblio = 'Отсутствуют источники или их заголовок задан неверно' 
        else:
            temp = ''
            flg = 1
            number = 1
            for i in self.document.paragraphs[maxi + 1:-1]:
                if 'Приложение' in i.text:
                    flg = 0
                if flg:
                    if re.sub(r"^\s+", "", i.text) != '':
                        temp = temp + number.__str__() + '. ' + i.text + '\n'

            self.biblio = [checkBiblio(i) for i in temp.split('\n')]



def findParagraph(text, document):
    num_par = 0
    for paragraph in document.paragraphs:
        if (text.lower() in paragraph.text.strip().lower()):
            if ((len(paragraph.text.strip()) - len(text)) == 0):
                return num_par
        num_par += 1
    return -1


if __name__ == '__main__':
    import glob
    checkDocuments = []
    for i in glob.glob('/home/danila/Test/docx/*.docx'):
        checkDocuments.append(checkDocument(i))

    for i in checkDocuments:
        print(i.way)
        print(i.studentInfo)
        print(i.errorsTitle)
        print(i.biblio)
        print()
